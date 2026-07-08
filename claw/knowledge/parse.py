"""Extract plain text from uploaded knowledge documents.

Supports pdf, docx, html, markdown, and plain text. Parsing libraries are
imported lazily so a missing optional dep degrades to a clear error instead of
crashing import. The caller handles the returned error string.
"""

from __future__ import annotations

from html.parser import HTMLParser


class _TextExtractor(HTMLParser):
    """Collect visible text from HTML, skipping script/style/head noise."""

    _SKIP = {"script", "style", "head", "meta", "link", "noscript"}

    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in self._SKIP:
            self._skip_depth += 1
        elif tag in ("p", "br", "div", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6"):
            self._parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP and self._skip_depth > 0:
            self._skip_depth -= 1

    def handle_data(self, data: str) -> None:
        if self._skip_depth == 0 and data.strip():
            self._parts.append(data)

    def text(self) -> str:
        return "".join(self._parts)


def _clean(text: str) -> str:
    # Collapse runs of blank lines / trailing whitespace so chunks stay tidy.
    lines = [ln.rstrip() for ln in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    out: list[str] = []
    blanks = 0
    for ln in lines:
        if ln.strip():
            blanks = 0
            out.append(ln)
        else:
            blanks += 1
            if blanks <= 1:
                out.append("")
    return "\n".join(out).strip()


def extract_text(data: bytes, filename: str, mime: str) -> tuple[str, str]:
    """Return (text, error). On success error is ""; on failure text is ""."""
    name = (filename or "").lower()
    ext = name.rsplit(".", 1)[-1] if "." in name else ""

    try:
        if ext == "pdf" or mime == "application/pdf":
            import io

            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            pages = [(p.extract_text() or "") for p in reader.pages]
            return _clean("\n\n".join(pages)), ""

        if ext == "docx" or "wordprocessingml" in mime:
            import io

            import docx

            document = docx.Document(io.BytesIO(data))
            paras = [p.text for p in document.paragraphs]
            # Include table cell text too — tables often hold the real content.
            for table in document.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paras.append(" | ".join(cells))
            return _clean("\n".join(paras)), ""

        if ext in ("html", "htm") or mime in ("text/html", "application/xhtml+xml"):
            parser = _TextExtractor()
            parser.feed(data.decode("utf-8", "replace"))
            return _clean(parser.text()), ""

        if ext in ("md", "markdown", "txt", "text", "csv", "log") or mime.startswith("text/"):
            return _clean(data.decode("utf-8", "replace")), ""

        # Fall back to a best-effort UTF-8 decode for unknown types.
        text = _clean(data.decode("utf-8", "replace"))
        if text:
            return text, ""
        return "", f"unsupported file type: {filename or mime or 'unknown'}"
    except Exception as exc:  # noqa: BLE001 - surface a clean message to the user
        return "", f"could not read {filename}: {exc}"


def chunk_text(text: str, *, size: int = 900, overlap: int = 150) -> list[str]:
    """Split text into ~`size`-char chunks on paragraph boundaries with overlap,
    so a retrieved chunk carries enough surrounding context to answer from."""
    text = text.strip()
    if not text:
        return []
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    current = ""
    for para in paragraphs:
        if len(para) > size:
            # Flush what we have, then hard-split the oversized paragraph.
            if current:
                chunks.append(current)
                current = ""
            for i in range(0, len(para), size - overlap):
                chunks.append(para[i : i + size])
            continue
        if current and len(current) + len(para) + 2 > size:
            chunks.append(current)
            # Carry a tail of the previous chunk for continuity.
            tail = current[-overlap:] if overlap else ""
            current = (tail + "\n\n" + para).strip()
        else:
            current = (current + "\n\n" + para).strip() if current else para
    if current:
        chunks.append(current)
    return chunks


# A parsed slice of a document: (page, text). `page` is the 1-based PDF page
# number, or None for formats without pages (docx/html/txt). Used to attribute
# retrieved chunks back to a page for citations.
Segment = tuple[int | None, str]


def _is_pdf(filename: str, mime: str) -> bool:
    name = (filename or "").lower()
    ext = name.rsplit(".", 1)[-1] if "." in name else ""
    return ext == "pdf" or mime == "application/pdf"


def _pdf_segments(source) -> list[Segment]:
    """Per-page (page, text) segments from a PDF. `source` is a path str or a
    binary stream — pypdf accepts either, so the bytes and file paths share one
    extraction path."""
    from pypdf import PdfReader

    reader = PdfReader(source)
    segments: list[Segment] = []
    for i, page in enumerate(reader.pages):
        cleaned = _clean(page.extract_text() or "")
        if cleaned:
            segments.append((i + 1, cleaned))
    return segments


def extract_segments(data: bytes, filename: str, mime: str) -> tuple[list[Segment], str]:
    """Like extract_text, but returns per-page segments for PDFs so chunks can
    be attributed to a page. Returns (segments, error); on failure segments=[].

    Non-paged formats (docx/html/txt/…) return a single (None, text) segment,
    so downstream chunking is byte-for-byte identical to the pre-page behavior.
    """
    try:
        if _is_pdf(filename, mime):
            import io

            return _pdf_segments(io.BytesIO(data)), ""
        # Every other format has no intrinsic pages — one segment, page=None.
        text, err = extract_text(data, filename, mime)
        if err:
            return [], err
        return ([(None, text)] if text else []), ""
    except Exception as exc:  # noqa: BLE001 - surface a clean message to the user
        return [], f"could not read {filename}: {exc}"


def _segments_char_count(segments: list[Segment]) -> int:
    return sum(len(t) for _, t in segments)


def ocr_available() -> bool:
    """True when the `ocrmypdf` CLI is on PATH (it also needs tesseract +
    ghostscript, which ocrmypdf checks at run time)."""
    import shutil

    return shutil.which("ocrmypdf") is not None


def _ocr_pdf_to_segments(path: str, *, timeout: int) -> tuple[list[Segment], str]:
    """Run ocrmypdf to add a text layer, then re-extract per-page text.
    Returns (segments, error)."""
    import subprocess
    import tempfile

    out_fd, out_path = tempfile.mkstemp(suffix=".pdf")
    import os

    os.close(out_fd)
    try:
        # --skip-text: only OCR pages that lack text (safe for mixed docs);
        # --optimize 0: skip slow image optimization we don't need.
        subprocess.run(
            ["ocrmypdf", "--skip-text", "--optimize", "0", path, out_path],
            check=True,
            capture_output=True,
            timeout=timeout,
        )
        return _pdf_segments(out_path), ""
    except subprocess.TimeoutExpired:
        return [], "OCR timed out"
    except subprocess.CalledProcessError as exc:
        detail = (exc.stderr or b"").decode("utf-8", "replace")[:300]
        return [], f"OCR failed: {detail}"
    except Exception as exc:  # noqa: BLE001
        return [], f"OCR failed: {exc}"
    finally:
        try:
            os.remove(out_path)
        except OSError:
            pass


def extract_segments_path(path: str, filename: str, mime: str) -> tuple[list[Segment], str]:
    """Path-based variant of extract_segments — reads PDFs/DOCX straight from the
    file (memory-light for large uploads) instead of a full in-memory bytes copy."""
    name = (filename or "").lower()
    ext = name.rsplit(".", 1)[-1] if "." in name else ""
    try:
        if _is_pdf(filename, mime):
            return _pdf_segments(path), ""
        if ext == "docx" or "wordprocessingml" in mime:
            import docx

            document = docx.Document(path)
            paras = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paras.append(" | ".join(cells))
            text = _clean("\n".join(paras))
            return ([(None, text)] if text else []), ""
        # Text-ish formats are small; read + reuse the bytes decoder.
        from pathlib import Path

        data = Path(path).read_bytes()
        text, err = extract_text(data, filename, mime)
        if err:
            return [], err
        return ([(None, text)] if text else []), ""
    except Exception as exc:  # noqa: BLE001
        return [], f"could not read {filename}: {exc}"


def chunk_segments(segments: list[Segment], *, size: int = 900, overlap: int = 150) -> list[Segment]:
    """Chunk each segment independently, tagging every chunk with its page.

    A PDF's chunks therefore never span two pages (so page attribution is exact);
    a single-segment (page=None) document chunks exactly as `chunk_text` did before.
    """
    out: list[Segment] = []
    for page, text in segments:
        for chunk in chunk_text(text, size=size, overlap=overlap):
            out.append((page, chunk))
    return out


def parse_document(
    data: bytes, filename: str, mime: str, *, size: int = 900, overlap: int = 150
) -> tuple[str, list[Segment], str]:
    """One CPU-bound pass: extract + chunk. Meant to run in a worker thread so it
    never blocks the event loop. Returns (full_text, chunk_records, error)."""
    segments, err = extract_segments(data, filename, mime)
    if err:
        return "", [], err
    full_text = _clean("\n\n".join(t for _, t in segments))
    chunk_records = chunk_segments(segments, size=size, overlap=overlap)
    return full_text, chunk_records, ""


def parse_document_path(
    path: str,
    filename: str,
    mime: str,
    *,
    ocr_enabled: bool = False,
    ocr_min_chars: int = 20,
    ocr_timeout: int = 600,
    size: int = 900,
    overlap: int = 150,
) -> tuple[str, list[Segment], str]:
    """Extract + chunk a document read from disk (memory-light for large files),
    with an optional OCR fallback for scanned PDFs. Run this in a worker thread.

    Returns (full_text, chunk_records, error). A scanned PDF with OCR disabled or
    unavailable returns a clear, actionable error instead of a silent empty doc.
    """
    segments, err = extract_segments_path(path, filename, mime)
    if err:
        return "", [], err

    if _is_pdf(filename, mime) and _segments_char_count(segments) < ocr_min_chars:
        # Little/no embedded text ⇒ almost certainly a scan.
        if ocr_enabled and ocr_available():
            ocr_segments, ocr_err = _ocr_pdf_to_segments(path, timeout=ocr_timeout)
            if ocr_err:
                return "", [], ocr_err
            segments = ocr_segments
        elif not segments:
            hint = (
                "enable OCR (CLAW_KNOWLEDGE__OCR_ENABLED=true and install ocrmypdf)"
                if not ocr_enabled
                else "OCR is enabled but the ocrmypdf tool isn't installed"
            )
            return "", [], f"no readable text found in {filename} — looks scanned; {hint}"

    full_text = _clean("\n\n".join(t for _, t in segments))
    chunk_records = chunk_segments(segments, size=size, overlap=overlap)
    return full_text, chunk_records, ""
